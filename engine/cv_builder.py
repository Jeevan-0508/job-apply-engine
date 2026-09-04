"""
Builds a tailored CV from the structured profile in config/profile.py,
reordered per-role against a specific JD's weighted skill signal, as .docx
and .pdf.

Two layouts share one code path:

- "international": headings and dates in English, no personal data beyond
  contact details. What Amazon and most international ATS expect.
- "german": a tabellarischer Lebenslauf -- German headings, a Persoenliche
  Daten block, MM/YYYY periods with "heute" for a current role, and a place,
  date and signature line at the end.

Both stay strictly single-column with no tables. A two-column CV is the
classic German look, but parsers interleave the columns into nonsense, so the
tabular feel is carried by the period sitting on the role line instead.

Text is set in Bitstream Vera, which ships with reportlab and so is present
on any machine that can install this app. That is not cosmetic: reportlab's
built-in Helvetica has no Unicode map for U+2022, so every bullet in a
Helvetica PDF extracts as "(cid:127)" and corrupts the line it sits on --
see engine/ats_check.py, which measures exactly that.

Falls back to a flat "top matched lines" mode if the profile hasn't been
filled in yet -- see build_from_flat_lines().
"""
import os
import re
from datetime import date

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor
from reportlab.lib.colors import HexColor
from reportlab.lib.enums import TA_RIGHT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
import pdfplumber
from reportlab.platypus import (HRFlowable, KeepTogether, Paragraph,
                                SimpleDocTemplate, Spacer)

import reportlab

from engine.jd_analyzer import match_profile, profile_corpus

BULLET = "\u2022"
ACCENT = HexColor("#1F3864")

LABELS = {
    "international": {
        "doc_title": None,
        "personal": None,
        "summary": "Professional Summary",
        "skills": "Core Skills",
        "experience": "Professional Experience",
        "education": "Education",
        "certifications": "Certifications",
        "languages": "Languages",
        "present": "Present",
        "born": "Date of birth",
        "nationality": "Nationality",
        "address": "Address",
        "phone": "Phone",
        "email": "Email",
    },
    "german": {
        "doc_title": "Lebenslauf",
        "personal": "Pers\u00f6nliche Daten",
        "summary": "Kurzprofil",
        "skills": "Kenntnisse",
        "experience": "Berufserfahrung",
        "education": "Ausbildung",
        "certifications": "Zertifikate und Weiterbildungen",
        "languages": "Sprachen",
        "present": "heute",
        "born": "Geburtsdatum",
        "nationality": "Staatsangeh\u00f6rigkeit",
        "address": "Adresse",
        "phone": "Telefon",
        "email": "E-Mail",
    },
}

_MONTH_YEAR = re.compile(r"^(0[1-9]|1[0-2])[./](\d{4})$")
_YEAR = re.compile(r"^\d{4}$")
_PRESENT = {"present", "heute", "current", "now", "today", "ongoing", ""}


def register_fonts():
    """Register Bitstream Vera and return the family names to use.

    Falls back to Helvetica if reportlab's bundled fonts are missing, which
    costs Unicode bullets -- ats_check will report it rather than hide it.
    """
    fonts_dir = os.path.join(os.path.dirname(reportlab.__file__), "fonts")
    faces = {"regular": "Vera.ttf", "bold": "VeraBd.ttf",
             "italic": "VeraIt.ttf", "bolditalic": "VeraBI.ttf"}
    names = {"regular": "CVSans", "bold": "CVSans-Bold",
             "italic": "CVSans-Italic", "bolditalic": "CVSans-BoldItalic"}
    try:
        for key, filename in faces.items():
            path = os.path.join(fonts_dir, filename)
            if not os.path.exists(path):
                raise FileNotFoundError(path)
            if names[key] not in pdfmetrics.getRegisteredFontNames():
                pdfmetrics.registerFont(TTFont(names[key], path))
        pdfmetrics.registerFontFamily(
            names["regular"], normal=names["regular"], bold=names["bold"],
            italic=names["italic"], boldItalic=names["bolditalic"])
        return names
    except Exception:
        return {"regular": "Helvetica", "bold": "Helvetica-Bold",
                "italic": "Helvetica-Oblique", "bolditalic": "Helvetica-BoldOblique"}


def _score_bullet(bullet, jd_signal):
    b = bullet.lower()
    return sum(weight for skill, weight in jd_signal.items() if skill in b)


def tailor_profile(profile, jd_signal, top_n_per_role=6):
    """Reorders bullets within each role by JD relevance. Never drops roles
    or fabricates content -- only reorders and trims to top_n per role."""
    tailored_experience = []
    for role in profile["experience"]:
        scored = sorted(
            role["bullets"],
            key=lambda b: _score_bullet(b, jd_signal),
            reverse=True,
        )
        tailored_experience.append({**role, "bullets": scored[:top_n_per_role]})

    matched_skills, gap_skills = match_profile(
        profile["skills"], jd_signal, profile_corpus(profile)
    )
    other_skills = [s for s in profile["skills"] if s not in matched_skills]

    return {
        **profile,
        "experience": tailored_experience,
        "skills": matched_skills + other_skills,  # matched ones surface first, nothing removed
        "gap_skills": gap_skills,  # shown in the app, never stuffed into the CV
    }


def _fmt_endpoint(value, labels):
    """Normalise one end of a period without inventing precision."""
    raw = (value or "").strip()
    if raw.lower() in _PRESENT:
        return labels["present"]
    if _MONTH_YEAR.match(raw):
        month, year = _MONTH_YEAR.match(raw).groups()
        return f"{month}/{year}"
    return raw


def period(role, labels):
    start = _fmt_endpoint(role.get("start"), labels)
    end = _fmt_endpoint(role.get("end"), labels)
    if start and end:
        return f"{start} \u2013 {end}"
    return start or end


def missing_month_precision(profile):
    """Periods given as a bare year, which a German CV is expected to avoid."""
    vague = []
    for role in profile.get("experience", []) or []:
        for key in ("start", "end"):
            raw = (role.get(key) or "").strip()
            if _YEAR.match(raw):
                vague.append(f"{role.get('role','?')} ({role.get('company','?')}) {key}={raw}")
    return vague


def _headline(profile):
    """First clause of the title only -- the rest is keyword soup for the
    skills section, not a job title, and reads badly under a name."""
    title = (profile.get("title") or "").strip()
    return title.split("|")[0].strip(" \u2014-")


def _contact_lines(profile, labels, layout):
    """Contact details on their own short lines.

    One long pipe-delimited run is the single most common way a parser mangles
    a header: it reads the whole thing as the job title.
    """
    if layout == "german":
        rows = [
            (labels["address"], profile.get("address") or profile.get("location")),
            (labels["phone"], profile.get("phone")),
            (labels["email"], profile.get("email")),
            (labels["born"], profile.get("dob")),
            (labels["nationality"], profile.get("nationality")),
            ("LinkedIn", profile.get("linkedin")),
            ("GitHub", profile.get("github")),
        ]
        return [f"{label}: {value}" for label, value in rows if value]

    primary = " \u00b7 ".join(filter(None, [
        profile.get("email"), profile.get("phone"), profile.get("location")]))
    secondary = " \u00b7 ".join(filter(None, [
        profile.get("linkedin"), profile.get("github")]))
    return [line for line in (primary, secondary) if line]


def _pdf_styles(fonts, scale=1.0):
    """Style set at a given density. scale<1 tightens type and spacing so a CV
    that spills a few lines onto a second page can be pulled back onto one."""
    base = getSampleStyleSheet()
    body = ParagraphStyle("CVBody", parent=base["BodyText"], fontName=fonts["regular"],
                          fontSize=9.5 * scale, leading=13 * scale, spaceAfter=3 * scale)
    return {
        "name": ParagraphStyle("CVName", parent=body, fontName=fonts["bold"],
                               fontSize=19 * scale, leading=22 * scale, spaceAfter=1, textColor=HexColor("#111111")),
        "doc_title": ParagraphStyle("CVDocTitle", parent=body, fontName=fonts["bold"],
                                    fontSize=11 * scale, leading=13 * scale, spaceAfter=8 * scale, textColor=ACCENT),
        "headline": ParagraphStyle("CVHeadline", parent=body, fontName=fonts["regular"],
                                   fontSize=10.5 * scale, leading=13 * scale, spaceAfter=4 * scale, textColor=ACCENT),
        "contact": ParagraphStyle("CVContact", parent=body, fontSize=9 * scale, leading=12 * scale, spaceAfter=1),
        # keepWithNext binds a heading to the line after it. KeepTogether on the
        # heading plus a whole role instead pushes a tall role to the next page
        # and leaves half of page 1 blank.
        "h1": ParagraphStyle("CVH1", parent=body, fontName=fonts["bold"], fontSize=11 * scale,
                             leading=13 * scale, spaceBefore=10 * scale,
                             spaceAfter=2, textColor=ACCENT,
                             keepWithNext=True),
        "role": ParagraphStyle("CVRole", parent=body, fontName=fonts["bold"], spaceAfter=0),
        "meta": ParagraphStyle("CVMeta", parent=body, fontName=fonts["italic"],
                               fontSize=9 * scale, textColor=HexColor("#444444"), spaceAfter=2),
        # The bullet is part of the text with a hanging indent rather than
        # reportlab's bulletText. bulletText is drawn as a separate text run, so
        # some parsers emit it as its own line and every bullet becomes an
        # orphan "\u2022" above the sentence it belongs to.
        "bullet": ParagraphStyle("CVBullet", parent=body, leftIndent=12,
                                 firstLineIndent=-12, spaceAfter=1.5 * scale),
        "body": body,
        "sign": ParagraphStyle("CVSign", parent=body, spaceBefore=16),
    }


def _rule():
    return HRFlowable(width="100%", thickness=0.6, color=ACCENT,
                      spaceBefore=1, spaceAfter=5)


def _page_count(path):
    try:
        with pdfplumber.open(path) as pdf:
            return len(pdf.pages)
    except Exception:
        return 1


def build_cv_pdf(tailored, out_path, layout="international", fit_pages=1):
    """Render the CV, tightening density until it fits fit_pages if it can.

    A CV that spills three lines onto a second page reads as padded, and the
    orphan page is usually the first thing a recruiter notices. Rather than
    silently dropping content, the same content is re-rendered denser and the
    result is measured; if it still will not fit, the honest longer version is
    kept.
    """
    for scale, inline_lists in ((1.0, False), (0.95, False), (0.92, True), (0.88, True), (0.84, True)):
        _render_pdf(tailored, out_path, layout, scale, inline_lists)
        if _page_count(out_path) <= fit_pages:
            break
    return out_path


def _render_pdf(tailored, out_path, layout, scale=1.0, inline_lists=False):
    fonts = register_fonts()
    labels = LABELS.get(layout, LABELS["international"])
    st = _pdf_styles(fonts, scale)

    doc = SimpleDocTemplate(
        out_path, pagesize=A4, topMargin=1.4 * cm, bottomMargin=1.3 * cm,
        leftMargin=1.9 * cm, rightMargin=1.9 * cm,
        title=f"{tailored.get('name','CV')} - CV", author=tailored.get("name", ""),
        subject=tailored.get("title", ""),
    )

    flow = []
    if labels["doc_title"]:
        flow.append(Paragraph(labels["doc_title"], st["doc_title"]))
    flow.append(Paragraph(tailored.get("name", ""), st["name"]))

    headline = _headline(tailored)
    if headline:
        flow.append(Paragraph(headline, st["headline"]))

    if labels["personal"]:
        flow.append(_rule())
        flow.append(Paragraph(labels["personal"], st["h1"]))
    for line in _contact_lines(tailored, labels, layout):
        flow.append(Paragraph(line, st["contact"]))

    def section(title, blocks):
        if not blocks:
            return
        flow.append(Paragraph(title, st["h1"]))
        flow.append(_rule())
        flow.extend(blocks)

    if tailored.get("summary"):
        section(labels["summary"], [Paragraph(tailored["summary"], st["body"])])

    if tailored.get("skills"):
        section(labels["skills"], [Paragraph(" \u00b7 ".join(tailored["skills"]), st["body"])])

    exp_blocks = []
    for role in tailored.get("experience", []) or []:
        block = [Paragraph(f"{role.get('role','')} \u2013 {role.get('company','')}", st["role"])]
        meta = " | ".join(filter(None, [period(role, labels), role.get("location")]))
        if meta:
            block.append(Paragraph(meta, st["meta"]))
        for bullet in role.get("bullets", []) or []:
            block.append(Paragraph(f"{BULLET} {bullet}", st["bullet"]))
        block.append(Spacer(1, 5))
        # a role's heading, period and first bullet must not split across pages;
        # the remaining bullets may flow
        exp_blocks.append(KeepTogether(block[:3]))
        exp_blocks.extend(block[3:])
    section(labels["experience"], exp_blocks)

    edu_blocks = []
    for ed in tailored.get("education", []) or []:
        edu_blocks.append(Paragraph(
            f"{ed.get('degree','')} \u2013 {ed.get('institution','')} ({ed.get('year','')})", st["body"]))
    section(labels["education"], edu_blocks)

    certs = tailored.get("certifications", []) or []
    if inline_lists and certs:
        cert_blocks = [Paragraph(" \u00b7 ".join(certs), st["body"])]
    else:
        cert_blocks = [Paragraph(f"{BULLET} {c}", st["bullet"]) for c in certs]
    section(labels["certifications"], cert_blocks)

    langs = tailored.get("languages", []) or []
    if langs:
        section(labels["languages"], [Paragraph(
            " \u00b7 ".join(f"{l['name']} ({l['level']})" for l in langs), st["body"])])

    if layout == "german":
        city = (tailored.get("signature_city")
                or (tailored.get("location") or "").split(",")[0].strip())
        stamp = f"{city}, {date.today().strftime('%d.%m.%Y')}" if city else date.today().strftime("%d.%m.%Y")
        flow.append(KeepTogether([
            Paragraph(stamp, st["sign"]),
            Spacer(1, 14),
            Paragraph(tailored.get("name", ""), st["body"]),
        ]))

    doc.build(flow)
    return out_path


def _docx_heading(doc, text, size=12):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    return p


def build_cv_docx(tailored, out_path, layout="international"):
    labels = LABELS.get(layout, LABELS["international"])
    doc = Document()

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10)

    if labels["doc_title"]:
        t = doc.add_paragraph()
        run = t.add_run(labels["doc_title"])
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)

    name = doc.add_paragraph()
    name_run = name.add_run(tailored.get("name", ""))
    name_run.bold = True
    name_run.font.size = Pt(19)
    name_run.font.color.rgb = RGBColor(0x11, 0x11, 0x11)
    name.paragraph_format.space_after = Pt(0)

    headline = _headline(tailored)
    if headline:
        h = doc.add_paragraph()
        h_run = h.add_run(headline)
        h_run.font.size = Pt(10.5)
        h_run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
        h.paragraph_format.space_after = Pt(2)

    if labels["personal"]:
        _docx_heading(doc, labels["personal"])
    for line in _contact_lines(tailored, labels, layout):
        p = doc.add_paragraph()
        p.add_run(line).font.size = Pt(9)
        p.paragraph_format.space_after = Pt(0)

    if tailored.get("summary"):
        _docx_heading(doc, labels["summary"])
        doc.add_paragraph(tailored["summary"])

    if tailored.get("skills"):
        _docx_heading(doc, labels["skills"])
        doc.add_paragraph(" \u00b7 ".join(tailored["skills"]))

    if tailored.get("experience"):
        _docx_heading(doc, labels["experience"])
        for role in tailored["experience"]:
            p = doc.add_paragraph()
            p.add_run(f"{role.get('role','')} \u2013 {role.get('company','')}").bold = True
            p.paragraph_format.space_after = Pt(0)
            meta = " | ".join(filter(None, [period(role, labels), role.get("location")]))
            if meta:
                m = doc.add_paragraph()
                m_run = m.add_run(meta)
                m_run.italic = True
                m_run.font.size = Pt(9)
                m.paragraph_format.space_after = Pt(2)
            for bullet in role.get("bullets", []) or []:
                doc.add_paragraph(bullet, style="List Bullet")

    if tailored.get("education"):
        _docx_heading(doc, labels["education"])
        for ed in tailored["education"]:
            doc.add_paragraph(
                f"{ed.get('degree','')} \u2013 {ed.get('institution','')} ({ed.get('year','')})")

    if tailored.get("certifications"):
        _docx_heading(doc, labels["certifications"])
        for c in tailored["certifications"]:
            doc.add_paragraph(c, style="List Bullet")

    if tailored.get("languages"):
        _docx_heading(doc, labels["languages"])
        doc.add_paragraph(" \u00b7 ".join(
            f"{l['name']} ({l['level']})" for l in tailored["languages"]))

    if layout == "german":
        city = (tailored.get("signature_city")
                or (tailored.get("location") or "").split(",")[0].strip())
        stamp = f"{city}, {date.today().strftime('%d.%m.%Y')}" if city else date.today().strftime("%d.%m.%Y")
        s = doc.add_paragraph(stamp)
        s.paragraph_format.space_before = Pt(16)
        doc.add_paragraph("")
        doc.add_paragraph(tailored.get("name", ""))

    doc.save(out_path)
    return out_path


def build_from_flat_lines(lines, out_path_docx, out_path_pdf, header="TAILORED RESUME - DRAFT"):
    """Fallback mode for when config/profile.py is still unfilled: takes the
    flat JD-ranked lines from engine/tailor.py and produces simple docx+pdf."""
    doc = Document()
    doc.add_heading(header, level=1)
    for line in lines:
        doc.add_paragraph(line, style="List Bullet")
    doc.save(out_path_docx)

    fonts = register_fonts()
    st = _pdf_styles(fonts)
    pdf = SimpleDocTemplate(out_path_pdf, pagesize=A4)
    flow = [Paragraph(header, st["h1"])]
    for line in lines:
        flow.append(Paragraph(f"{BULLET} {line}", st["bullet"]))
    pdf.build(flow)

    return out_path_docx, out_path_pdf
