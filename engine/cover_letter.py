"""
Generates a tailored cover letter from the structured profile plus a job's
matched skill signal, in English or as a German Anschreiben.

Honest by construction: it only ever cites skills and STAR proof points that
exist in the profile -- nothing invented.

The German version is not a translated English letter. A German Anschreiben has
a layout recruiters actively look for: sender block, recipient, right-aligned
place and date, a bold "Bewerbung als ..." subject line, a formal salutation
and "Mit freundlichen Gruessen". The sentences this module writes are authored
in German rather than machine-translated. Text taken from the profile -- STAR
examples, the motivation paragraph -- is inserted verbatim in whatever language
it was written in, and the app says so instead of pretending otherwise.
"""
from datetime import date

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from reportlab.lib.enums import TA_RIGHT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

from engine.cv_builder import register_fonts
from engine.jd_analyzer import match_profile, profile_corpus


def _pick_proof_point(profile, jd_signal):
    """Pick the STAR example with the most overlap against this JD's skills."""
    best, best_score = None, -1
    for ex in profile.get("star_examples", []) or []:
        score = sum(1 for tag in ex.get("skills_tags", []) if tag in jd_signal)
        if score > best_score:
            best, best_score = ex, score
    return best


def _top_skills(profile, jd_signal):
    """Skills to claim in the letter -- only ones the profile actually evidences.

    This must never fall back to the job's own requirements. Doing so produced
    letters opening "my experience in loss prevention, fraud investigation"
    naming precisely the skills the profile did not have, which is the one thing
    this module promises not to do. If nothing matches, the profile's own
    strongest skills are named instead and the overlap claim is dropped.
    """
    matched, _ = match_profile(profile.get("skills", []), jd_signal, profile_corpus(profile))
    if matched:
        return matched[:4], True
    return (profile.get("skills") or [])[:3], False


def build_cover_letter_text(profile, jd_signal, company, role, language="en"):
    """Body paragraphs only, in reading order. Layout is applied by the writers."""
    claimed, overlaps = _top_skills(profile, jd_signal)
    skills = ", ".join(claimed) or profile.get("title", "")
    proof = _pick_proof_point(profile, jd_signal)
    proof_text = ""
    if proof and proof.get("title") != "[FILL IN]":
        proof_text = " ".join(filter(None, [
            proof.get("situation"), proof.get("task"), proof.get("action")]))
        if proof.get("result"):
            proof_text += f" Result: {proof['result']}"

    motivation = profile.get("why_germany", "")

    if language == "de":
        paragraphs = [
            "Sehr geehrte Damen und Herren,",
            (f"mit gro\u00dfem Interesse bewerbe ich mich auf die Position als {role}. "
             f"Meine Erfahrung in {skills} entspricht dem von Ihnen beschriebenen "
             f"Anforderungsprofil, und ich m\u00f6chte diese gerne in Ihr Team einbringen."
             if overlaps else
             f"mit gro\u00dfem Interesse bewerbe ich mich auf die Position als {role}. "
             f"Mein Schwerpunkt liegt in {skills}, und ich m\u00f6chte diese Erfahrung "
             f"gerne in Ihr Team einbringen."),
            (f"Ein Beispiel aus meiner Praxis: {proof_text}" if proof_text else ""),
            motivation,
            ("\u00dcber die Gelegenheit zu einem pers\u00f6nlichen Gespr\u00e4ch w\u00fcrde ich mich sehr freuen."),
            "Mit freundlichen Gr\u00fc\u00dfen",
            profile.get("name", ""),
        ]
    else:
        paragraphs = [
            f"Dear Hiring Team at {company},",
            (f"I'm writing to apply for the {role} position. My background in {skills} "
             f"lines up closely with what you're looking for, and I'd welcome the chance "
             f"to bring that to your team."
             if overlaps else
             f"I'm writing to apply for the {role} position. My background is in {skills}, "
             f"and I'd welcome the chance to bring that experience to your team."),
            (f"In a recent example, {proof_text}" if proof_text else ""),
            motivation,
            "Thank you for your time and consideration -- I'd welcome the opportunity to discuss further.",
            "Best regards,",
            profile.get("name", ""),
        ]
    return [p for p in paragraphs if p]


def subject_line(company, role, language="en"):
    return f"Bewerbung als {role}" if language == "de" else f"Application: {role} at {company}"


def _stamp(profile, language):
    city = (profile.get("signature_city")
            or (profile.get("location") or "").split(",")[0].strip())
    if language == "de":
        today = date.today().strftime("%d.%m.%Y")
    else:
        today = date.today().strftime("%d %B %Y")
    return f"{city}, {today}" if city else today


def build_cover_letter_docx(profile, jd_signal, company, role, out_path, language="en"):
    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)

    for line in filter(None, [profile.get("name"), profile.get("address") or profile.get("location"),
                              profile.get("phone"), profile.get("email")]):
        p = doc.add_paragraph()
        p.add_run(line).font.size = Pt(10)
        p.paragraph_format.space_after = Pt(0)

    doc.add_paragraph("")
    recipient = doc.add_paragraph()
    recipient.add_run(company).bold = True
    recipient.paragraph_format.space_after = Pt(0)

    stamp = doc.add_paragraph(_stamp(profile, language))
    stamp.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    subject = doc.add_paragraph()
    subject.add_run(subject_line(company, role, language)).bold = True

    for para in build_cover_letter_text(profile, jd_signal, company, role, language):
        doc.add_paragraph(para)

    doc.save(out_path)
    return out_path


def build_cover_letter_pdf(profile, jd_signal, company, role, out_path, language="en"):
    fonts = register_fonts()
    base = getSampleStyleSheet()["BodyText"]
    body = ParagraphStyle("CLBody", parent=base, fontName=fonts["regular"],
                          fontSize=10.5, leading=15, spaceAfter=8)
    small = ParagraphStyle("CLSmall", parent=body, fontSize=9.5, leading=12, spaceAfter=0)
    right = ParagraphStyle("CLRight", parent=body, alignment=TA_RIGHT)
    bold = ParagraphStyle("CLBold", parent=body, fontName=fonts["bold"], spaceBefore=6)

    doc = SimpleDocTemplate(out_path, pagesize=A4, topMargin=2 * cm, bottomMargin=2 * cm,
                            leftMargin=2.2 * cm, rightMargin=2.2 * cm,
                            title=subject_line(company, role, language),
                            author=profile.get("name", ""))

    flow = []
    for line in filter(None, [profile.get("name"), profile.get("address") or profile.get("location"),
                              profile.get("phone"), profile.get("email")]):
        flow.append(Paragraph(line, small))
    flow.append(Spacer(1, 18))
    flow.append(Paragraph(company, bold))
    flow.append(Paragraph(_stamp(profile, language), right))
    flow.append(Spacer(1, 6))
    flow.append(Paragraph(subject_line(company, role, language), bold))
    flow.append(Spacer(1, 6))
    for para in build_cover_letter_text(profile, jd_signal, company, role, language):
        flow.append(Paragraph(para, body))

    doc.build(flow)
    return out_path
